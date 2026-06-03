"""Word 文档解析器

使用 python-docx 提取 DOCX 文本和表格的解析器实现。
包含 XXE 防护（defusedxml 预校验 DOCX 内 XML 文件）。
"""

from __future__ import annotations

import io
import logging
import os
import uuid
import zipfile
from datetime import UTC, datetime

from src.domain.ports.document_parser import DocumentParserPort
from src.domain.value_objects.parsed_document import (
    ParsedDocument,
    ParsedElement,
    ParsedPage,
    ParsedTable,
)
from src.infrastructure.document_parsing._limits import MAX_DOCX_BYTES

try:
    from defusedxml.ElementTree import parse as safe_xml_parse

    HAS_DEFUSEDXML = True
except ImportError:
    HAS_DEFUSEDXML = False

logger = logging.getLogger(__name__)


class WordParser(DocumentParserPort):
    """Word 文档解析器

    使用 python-docx.Document 提取文本和表格内容，支持：
    - 段落文本提取（含标题样式识别）
    - 表格行列结构提取
    - 旧版 DOC 格式拒绝
    - 文件大小上限保护（防御内嵌 OOXML 解压炸弹）
    """

    def parse(self, file_path: str, mime_type: str) -> ParsedDocument:
        """解析 DOCX 文件

        Args:
            file_path: 本地 DOCX 文件路径
            mime_type: MIME 类型（用于 DOC 格式拒绝）

        Returns:
            结构化解析结果
        """
        doc_id = str(uuid.uuid4())
        timestamp = datetime.now(UTC).isoformat()

        if mime_type == "application/msword":
            return ParsedDocument(
                document_id=doc_id,
                mime_type=mime_type,
                parse_status="failed",
                error_message="不支持旧版 DOC 格式，请转换为 DOCX",
                parse_timestamp=timestamp,
            )

        # 防御解压炸弹：DOCX 内嵌 OOXML 可塞 10GB+，解析前必须校验
        try:
            file_size = os.path.getsize(file_path)
        except OSError:
            logger.exception("DOCX 文件大小检查失败")
            return ParsedDocument(
                document_id=doc_id,
                mime_type=mime_type,
                parse_status="failed",
                error_message="无法访问文件，请检查文件路径或权限",
                parse_timestamp=timestamp,
            )
        if file_size > MAX_DOCX_BYTES:
            size_mb = file_size // (1024 * 1024)
            limit_mb = MAX_DOCX_BYTES // (1024 * 1024)
            return ParsedDocument(
                document_id=doc_id,
                mime_type=mime_type,
                parse_status="failed",
                error_message=(f"DOCX 文件大小 {size_mb}MB 超过 {limit_mb}MB 限制，可能为解压炸弹"),
                parse_timestamp=timestamp,
            )

        try:
            from docx import Document

            # XXE 防护：在 python-docx 解析前，用 defusedxml 预校验 DOCX 内所有 XML
            if HAS_DEFUSEDXML:
                try:
                    with zipfile.ZipFile(file_path, "r") as z:
                        for name in z.namelist():
                            if name.endswith(".xml"):
                                safe_xml_parse(io.BytesIO(z.read(name)))
                except Exception:
                    logger.exception("DOCX 包含不安全的 XML 内容")
                    return ParsedDocument(
                        document_id=doc_id,
                        mime_type=mime_type,
                        parse_status="failed",
                        error_message="DOCX 包含不安全的 XML 内容，解析已拒绝",
                        parse_timestamp=timestamp,
                    )

            with open(file_path, "rb") as f:
                doc = Document(f)

                # 空文档检测（AC-2 要求空 DOCX 返回解析失败）
                if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
                    return ParsedDocument(
                        document_id=doc_id,
                        mime_type=mime_type,
                        parse_status="failed",
                        error_message="DOCX 文档为空，未包含任何段落或表格",
                        parse_timestamp=timestamp,
                    )

                texts: list[ParsedElement] = []
                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        style_name = paragraph.style.name if paragraph.style else ""
                        texts.append(ParsedElement(content=text, metadata={"style": style_name}))

                tables: list[ParsedTable] = []
                for table in doc.tables:
                    rows = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        rows.append(row_data)
                    if rows:
                        tables.append(ParsedTable(rows=rows))

                page = ParsedPage(
                    page_number=1,
                    texts=texts,
                    tables=tables,
                    images=[],
                )

                return ParsedDocument(
                    document_id=doc_id,
                    mime_type=mime_type,
                    pages=[page],
                    parse_status="completed",
                    parse_timestamp=timestamp,
                )
        except Exception:
            # 安全：原始异常可能含文件路径，详细 traceback 记录到日志
            logger.exception("DOCX 文件解析失败")
            return ParsedDocument(
                document_id=doc_id,
                mime_type=mime_type,
                parse_status="failed",
                error_message="DOCX 解析失败，请检查文件是否损坏或重试",
                parse_timestamp=timestamp,
            )
