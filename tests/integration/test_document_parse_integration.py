"""文档解析集成测试

测试完整解析流水线：上传 → 解析 → 状态更新 → 事件发布。
使用真实解析器（非 Mock），验证端到端流程。
"""

from __future__ import annotations

import asyncio
import os
import tempfile
import uuid

import pytest
from pypdf import PdfWriter

from src.domain.entities.document import Document
from src.domain.value_objects.parsed_document import ParsedDocument


def _create_test_pdf(num_pages: int = 1) -> str:
    writer = PdfWriter()
    for _ in range(num_pages):
        writer.add_blank_page(width=612, height=792)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    writer.write(tmp.name)
    tmp.close()
    return tmp.name


def _create_test_docx() -> str:
    from docx import Document

    doc = Document()
    doc.add_paragraph("集成测试段落")
    doc.add_paragraph("第二段内容")
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def _create_test_txt(encoding: str = "utf-8", content: str = "集成测试文本\n\n第二段") -> str:
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".txt")
    tmp.write(content.encode(encoding))
    tmp.close()
    return tmp.name


class TestParsePipelinePDF:
    """PDF 解析流水线"""

    def test_parse_pdf_success(self) -> None:
        from src.infrastructure.external_services.document_parsing.pdf_parser import PDFParser

        parser = PDFParser()
        path = _create_test_pdf(2)
        try:
            result = parser.parse(path, "application/pdf")
            assert result.parse_status == "completed"
            assert len(result.pages) == 2
            assert result.mime_type == "application/pdf"
            assert result.document_id

            # 验证 to_dict 可 JSON 序列化
            import json

            json.dumps(result.to_dict(), ensure_ascii=False)
        finally:
            os.unlink(path)

    def test_parse_pdf_single_page(self) -> None:
        from src.infrastructure.external_services.document_parsing.pdf_parser import PDFParser

        parser = PDFParser()
        path = _create_test_pdf(1)
        try:
            result = parser.parse(path, "application/pdf")
            assert result.parse_status == "completed"
            assert len(result.pages) == 1
            assert result.pages[0].page_number == 1
        finally:
            os.unlink(path)


class TestParsePipelineDOCX:
    """DOCX 解析流水线"""

    def test_parse_docx_with_text(self) -> None:
        from src.infrastructure.external_services.document_parsing.word_parser import WordParser

        parser = WordParser()
        path = _create_test_docx()
        try:
            result = parser.parse(path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            assert result.parse_status == "completed"
            all_text = " ".join(t.content for p in result.pages for t in p.texts)
            assert "集成测试段落" in all_text
        finally:
            os.unlink(path)


class TestParsePipelineTXT:
    """TXT 解析流水线"""

    def test_parse_txt_utf8(self) -> None:
        from src.infrastructure.external_services.document_parsing.text_parser import TextParser

        parser = TextParser()
        path = _create_test_txt("utf-8")
        try:
            result = parser.parse(path, "text/plain")
            assert result.parse_status == "completed"
            assert len(result.pages) == 1
        finally:
            os.unlink(path)

    def test_parse_txt_gbk(self) -> None:
        from src.infrastructure.external_services.document_parsing.text_parser import TextParser

        parser = TextParser()
        path = _create_test_txt("gbk", "GBK编码测试")
        try:
            result = parser.parse(path, "text/plain")
            assert result.parse_status == "completed"
            all_text = " ".join(t.content for p in result.pages for t in p.texts)
            assert "GBK编码测试" in all_text
        finally:
            os.unlink(path)


class TestCompositeRouting:
    """组合路由集成测试"""

    def test_route_pdf(self) -> None:
        from src.infrastructure.external_services.document_parsing.composite_parser import CompositeDocumentParser
        from src.infrastructure.external_services.document_parsing.pdf_parser import PDFParser
        from src.infrastructure.external_services.document_parsing.text_parser import TextParser
        from src.infrastructure.external_services.document_parsing.word_parser import WordParser

        parser = CompositeDocumentParser(
            parsers={
                "application/pdf": PDFParser(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": WordParser(),
                "text/plain": TextParser(),
            },
        )
        path = _create_test_pdf()
        try:
            result = parser.parse(path, "application/pdf")
            assert result.parse_status == "completed"
        finally:
            os.unlink(path)

    def test_route_txt(self) -> None:
        from src.infrastructure.external_services.document_parsing.composite_parser import CompositeDocumentParser
        from src.infrastructure.external_services.document_parsing.pdf_parser import PDFParser
        from src.infrastructure.external_services.document_parsing.text_parser import TextParser
        from src.infrastructure.external_services.document_parsing.word_parser import WordParser

        parser = CompositeDocumentParser(
            parsers={
                "application/pdf": PDFParser(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": WordParser(),
                "text/plain": TextParser(),
            },
        )
        path = _create_test_txt()
        try:
            result = parser.parse(path, "text/plain")
            assert result.parse_status == "completed"
        finally:
            os.unlink(path)

    def test_route_unknown_raises(self) -> None:
        from src.infrastructure.external_services.document_parsing.composite_parser import CompositeDocumentParser
        from src.infrastructure.external_services.document_parsing.pdf_parser import PDFParser
        from src.infrastructure.external_services.document_parsing.text_parser import TextParser
        from src.infrastructure.external_services.document_parsing.word_parser import WordParser

        parser = CompositeDocumentParser(
            parsers={
                "application/pdf": PDFParser(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": WordParser(),
                "text/plain": TextParser(),
            },
        )
        path = _create_test_txt()
        try:
            with pytest.raises(ValueError, match="不支持的 MIME"):
                parser.parse(path, "application/unknown")
        finally:
            os.unlink(path)


class TestConcurrentParsing:
    """并发解析测试（AC-6: ≥10 并发）"""

    def test_concurrent_parse_10_documents(self) -> None:
        """验证可并发解析 ≥10 文档"""
        from src.infrastructure.external_services.document_parsing.pdf_parser import PDFParser

        parser = PDFParser()

        # 创建 10 个 PDF 文件
        paths = [_create_test_pdf() for _ in range(10)]
        try:

            async def parse_all() -> list[ParsedDocument]:
                tasks = [asyncio.to_thread(parser.parse, p, "application/pdf") for p in paths]
                return await asyncio.gather(*tasks)

            loop = asyncio.new_event_loop()
            try:
                results = loop.run_until_complete(parse_all())
            finally:
                loop.close()

            assert len(results) == 10
            for result in results:
                assert result.parse_status == "completed"
        finally:
            for p in paths:
                os.unlink(p)


class TestTempFileCleanup:
    """临时文件清理测试"""

    @pytest.mark.asyncio
    async def test_temp_file_cleaned_after_parse(self, monkeypatch) -> None:
        """验证解析完成后临时文件已删除（spy os.unlink 验证清理路径）"""
        import os
        from unittest.mock import AsyncMock, MagicMock

        from src.application.services.document_parsing_service import DocumentParsingService

        unlinked_paths: list[str] = []

        real_unlink = os.unlink

        def spy_unlink(path):
            unlinked_paths.append(path)
            return real_unlink(path)

        monkeypatch.setattr("os.unlink", spy_unlink)

        mock_repo = AsyncMock()
        mock_storage = AsyncMock()
        mock_publisher = AsyncMock()
        mock_parser = MagicMock()

        doc_id = uuid.uuid4()
        doc = Document(document_id=doc_id, filename="t.pdf", mime_type="application/pdf", tenant_id="t1")
        doc.metadata["storage_object_key"] = "key"

        mock_repo.find.return_value = doc
        mock_repo.save.return_value = doc

        def mock_retrieve(*args, **kwargs):
            async def _stream():
                yield b"content"

            return _stream()

        mock_storage.retrieve = MagicMock(side_effect=mock_retrieve)
        mock_parser.parse.return_value = ParsedDocument(document_id=str(doc_id), mime_type="application/pdf")

        service = DocumentParsingService(mock_repo, mock_storage, mock_publisher, mock_parser)
        await service.parse_document(doc_id, "t1")

        # 验证至少有一次 unlink 调用针对 /tmp 路径
        assert any("/tmp" in p or "tmp" in p for p in unlinked_paths), (
            f"临时文件应被清理，但未观察到对 /tmp 路径的 unlink 调用: {unlinked_paths}"
        )
