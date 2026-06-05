"""文档解析集成测试

测试完整解析流水线：
1. 解析器层：直接调用 PDF/DOCX/TXT 解析器
2. Service 层：通过 DocumentParsingService 编排完整流水线
   - 真实 PostgreSQL（schema 隔离）
   - 真实 MinIO（上传 → 下载 → 解析）
   - 真实 DualChannelEventBus（Redis + PostgreSQL Outbox → RabbitMQ）
   - 真实解析器
3. 并发解析测试（AC-6: ≥10 并发）
"""

from __future__ import annotations

import asyncio
import os
import tempfile
import uuid
from collections.abc import AsyncGenerator

import pytest
from pypdf import PdfWriter
from sqlalchemy import select as sa_select
from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession

from src.domain.entities.document import Document, ParseStatus
from src.domain.value_objects.parsed_document import ParsedDocument
from src.infrastructure.config.postgresql import PostgreSQLConfig
from src.infrastructure.storage.postgresql.models import OutboxModel
from src.infrastructure.storage.postgresql.postgresql_manager import PostgreSQLManager
from src.infrastructure.storage.postgresql.repository.document_repository import PostgreSQLDocumentRepository
from src.infrastructure.storage.postgresql.session_context import reset_session, set_session
from tests.environments import get_test_env

# ===================================================================
# Helpers
# ===================================================================


def _create_test_pdf(num_pages: int = 1) -> str:
    writer = PdfWriter()
    for _ in range(num_pages):
        writer.add_blank_page(width=612, height=792)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    writer.write(tmp.name)
    tmp.close()
    return tmp.name


def _create_test_pdf_with_text(text: str = "Integration Test Content") -> str:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    c = canvas.Canvas(tmp.name, pagesize=letter)
    c.drawString(72, 720, text)
    c.showPage()
    c.save()
    tmp.close()
    return tmp.name


def _create_test_docx() -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("集成测试段落")
    doc.add_paragraph("第二段内容")
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def _create_test_txt(encoding: str = "utf-8", content: str = "集成测试文本\n\n第二段") -> str:
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".txt")
    tmp.write(content.encode(encoding))
    tmp.close()
    return tmp.name


def _cleanup(path: str) -> None:
    """安全清理临时文件"""
    if path and os.path.exists(path):
        try:
            os.unlink(path)
        except OSError:
            pass


# ===================================================================
# PostgreSQL Schema Isolation Fixtures
# ===================================================================


@pytest.fixture
def test_schema() -> str:
    return f"test_sisys_{uuid.uuid4().hex[:8]}"


@pytest.fixture
def pg_config() -> PostgreSQLConfig:
    env = get_test_env()
    return PostgreSQLConfig(
        host=env.postgres.host,
        port=env.postgres.port,
        database=env.postgres.database,
        username=env.postgres.username,
        password=env.postgres.password,
        pool_size=5,
        max_overflow=10,
    )


@pytest.fixture
def db_engine(pg_config: PostgreSQLConfig) -> PostgreSQLManager:
    return PostgreSQLManager(pg_config)


@pytest.fixture
def ensure_schema(db_engine: PostgreSQLManager, pg_config: PostgreSQLConfig, test_schema: str):
    """创建测试专用 schema 及 documents 表"""
    sync_url = (
        f"postgresql+psycopg2://{pg_config.username}:{pg_config.password}"
        f"@{pg_config.host}:{pg_config.port}/{pg_config.database}"
    )
    from sqlalchemy import create_engine

    sync_engine = create_engine(sync_url)

    try:
        with sync_engine.connect() as conn:
            conn.execute(text("SELECT 1"))
    except Exception as e:
        sync_engine.dispose()
        pytest.skip(f"PostgreSQL not available: {e}")

    with sync_engine.connect() as conn:
        conn.execute(text(f'DROP SCHEMA IF EXISTS "{test_schema}" CASCADE'))
        conn.commit()

    with sync_engine.connect() as conn:
        conn.execute(text(f'CREATE SCHEMA "{test_schema}"'))
        conn.commit()

    from src.infrastructure.storage.postgresql.models import Base

    with sync_engine.connect() as conn:
        conn.execute(text(f'SET search_path TO "{test_schema}"'))
        Base.metadata.create_all(conn)
        conn.commit()

    sync_engine.dispose()

    yield test_schema

    sync_engine = create_engine(sync_url)
    try:
        with sync_engine.connect() as conn:
            conn.execute(text(f'DROP SCHEMA "{test_schema}" CASCADE'))
            conn.commit()
    except Exception:
        pass
    sync_engine.dispose()


@pytest.fixture
async def pg_session(db_engine: PostgreSQLManager, ensure_schema: str) -> AsyncGenerator[AsyncSession, None]:
    """创建带 schema 隔离的 PostgreSQL 会话"""
    async_engine = db_engine.get_async_engine()
    session = AsyncSession(async_engine, expire_on_commit=False)

    await session.execute(text(f'SET search_path TO "{ensure_schema}"'))

    async with session.begin_nested():
        yield session

    await session.rollback()
    await session.close()


@pytest.fixture
def repo(pg_session: AsyncSession):
    """注入 session 上下文的 DocumentRepository"""
    token = set_session(pg_session)
    repository = PostgreSQLDocumentRepository()
    yield repository
    reset_session(token)


# ===================================================================
# MinIO 真实服务 Fixtures
# ===================================================================


@pytest.fixture
def minio_tenant_id() -> str:
    """MinIO bucket 解析用的租户标识符"""
    return f"test-integration-{uuid.uuid4().hex[:8]}"


@pytest.fixture
def minio_storage(minio_tenant_id: str):
    """创建真实 MinIO 存储适配器，确保 bucket 存在，测试后清理"""
    from src.infrastructure.config.minio import MinIOConfig
    from src.infrastructure.storage.minio.bucket_manager import BucketManager
    from src.infrastructure.storage.minio.minio_adapter import MinIOAdapter
    from src.infrastructure.storage.minio.minio_repository import MinIORepository
    from src.infrastructure.storage.minio.object_operations import ObjectOperations
    from src.infrastructure.storage.minio.worm_lifecycle import WORMManager

    env = get_test_env()
    endpoint_parts = env.minio.endpoint.split(":")
    host = endpoint_parts[0]
    port = int(endpoint_parts[1]) if len(endpoint_parts) > 1 else 9000

    config = MinIOConfig(
        host=host,
        port=port,
        access_key=env.minio.access_key,
        secret_key=env.minio.secret_key,
        secure=env.minio.secure,
    )

    # 验证 MinIO 连通性
    try:
        from src.infrastructure.storage.minio.minio_manager import MinioManager

        mgr = MinioManager.from_config(config)
        if not mgr.health_check():
            pytest.skip("MinIO not available")
    except Exception as e:
        pytest.skip(f"MinIO not available: {e}")

    bucket_manager = BucketManager(config)
    object_ops = ObjectOperations(config)
    worm_mgr = WORMManager(config)

    repository = MinIORepository(
        bucket_manager=bucket_manager,
        object_operations=object_ops,
        worm_manager=worm_mgr,
        tenant_id=minio_tenant_id,
    )

    adapter = MinIOAdapter(repository)

    # 确保 raw-documents bucket 存在
    bucket_name = bucket_manager.build_bucket_name("raw-documents", minio_tenant_id)
    if not bucket_manager.bucket_exists(bucket_name):
        bucket_manager.create_bucket(bucket_name)

    yield adapter

    # 清理：删除测试 bucket 及其中所有对象
    _cleanup_minio_bucket(bucket_manager, bucket_name)


def _cleanup_minio_bucket(bucket_manager, bucket_name: str) -> None:
    """安全清理 MinIO bucket"""
    try:
        if bucket_manager.bucket_exists(bucket_name):
            bucket_manager.delete_bucket(bucket_name, force=True)
    except Exception:
        pass


def _upload_to_minio(adapter, bucket_type: str, object_key: str, file_path: str) -> None:
    """通过 MinIOAdapter 上传文件到 MinIO（同步包装，用于非 async 测试）"""
    import asyncio as _asyncio

    async def _upload():
        await adapter.store(bucket_type, object_key, file_path, "application/octet-stream")

    loop = _asyncio.new_event_loop()
    try:
        loop.run_until_complete(_upload())
    finally:
        loop.close()


async def _upload_to_minio_async(adapter, bucket_type: str, object_key: str, file_path: str) -> None:
    """通过 MinIOAdapter 上传文件到 MinIO（async 版本，用于 pytest-asyncio 测试）"""
    await adapter.store(bucket_type, object_key, file_path, "application/octet-stream")


# ===================================================================
# 真实 EventPublisher Fixture
# ===================================================================


@pytest.fixture
def event_bus():
    """真实 DualChannelEventBus（Redis + PostgreSQL Outbox → RabbitMQ）

    DocumentProcessed 路由到 RELIABLE 通道 → RabbitMQEventBus → OutboxRepository
    """
    from src.infrastructure.config.redis import RedisConfig as InfraRedisConfig
    from src.infrastructure.messaging.channel_router import ChannelRouter
    from src.infrastructure.messaging.dual_channel_event_bus import DualChannelEventBus
    from src.infrastructure.messaging.outbox.outbox_repository import PostgreSQLOutboxRepository
    from src.infrastructure.messaging.rabbitmq_event_bus import RabbitMQEventBus
    from src.infrastructure.messaging.redis_event_bus import RedisEventBus
    from src.infrastructure.messaging.redis_publisher import RedisEventPublisher
    from src.infrastructure.messaging.redis_subscriber import RedisEventSubscriber

    env = get_test_env()
    redis_config = InfraRedisConfig(
        host=env.redis.host,
        port=env.redis.port,
    )
    redis_publisher = RedisEventPublisher(redis_config)
    redis_subscriber = RedisEventSubscriber(redis_config)
    router = ChannelRouter(load_defaults=True)
    redis_bus = RedisEventBus(redis_publisher, redis_subscriber, router)
    rabbitmq_bus = RabbitMQEventBus(PostgreSQLOutboxRepository(), router)
    return DualChannelEventBus(redis_bus, rabbitmq_bus, router)


# ===================================================================
# 解析器层集成测试
# ===================================================================


class TestParsePipelinePDF:
    """PDF 解析流水线"""

    def test_parse_pdf_success(self) -> None:
        from src.infrastructure.document_parsing.pdf_parser import PDFParser

        parser = PDFParser()
        path = _create_test_pdf(2)
        try:
            result = parser.parse(path, "application/pdf")
            assert result.parse_status == "completed"
            assert len(result.pages) == 2
            assert result.mime_type == "application/pdf"
            assert result.document_id

            import json

            json.dumps(result.to_dict(), ensure_ascii=False)
        finally:
            os.unlink(path)

    def test_parse_pdf_single_page(self) -> None:
        from src.infrastructure.document_parsing.pdf_parser import PDFParser

        parser = PDFParser()
        path = _create_test_pdf(1)
        try:
            result = parser.parse(path, "application/pdf")
            assert result.parse_status == "completed"
            assert len(result.pages) == 1
            assert result.pages[0].page_number == 1
        finally:
            os.unlink(path)


class TestParsePipelineDOCX:
    """DOCX 解析流水线"""

    def test_parse_docx_with_text(self) -> None:
        from src.infrastructure.document_parsing.word_parser import WordParser

        parser = WordParser()
        path = _create_test_docx()
        try:
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            result = parser.parse(path, mime)
            assert result.parse_status == "completed"
            all_text = " ".join(t.content for p in result.pages for t in p.texts)
            assert "集成测试段落" in all_text
        finally:
            os.unlink(path)


class TestParsePipelineTXT:
    """TXT 解析流水线"""

    def test_parse_txt_utf8(self) -> None:
        from src.infrastructure.document_parsing.text_parser import TextParser

        parser = TextParser()
        path = _create_test_txt("utf-8")
        try:
            result = parser.parse(path, "text/plain")
            assert result.parse_status == "completed"
            assert len(result.pages) == 1
        finally:
            os.unlink(path)

    def test_parse_txt_gbk(self) -> None:
        from src.infrastructure.document_parsing.text_parser import TextParser

        parser = TextParser()
        path = _create_test_txt("gbk", "GBK编码测试")
        try:
            result = parser.parse(path, "text/plain")
            assert result.parse_status == "completed"
            all_text = " ".join(t.content for p in result.pages for t in p.texts)
            assert "GBK编码测试" in all_text
        finally:
            os.unlink(path)


class TestCompositeRouting:
    """组合路由集成测试"""

    def test_route_pdf(self) -> None:
        from src.infrastructure.document_parsing.composite_parser import CompositeDocumentParser
        from src.infrastructure.document_parsing.pdf_parser import PDFParser
        from src.infrastructure.document_parsing.text_parser import TextParser
        from src.infrastructure.document_parsing.word_parser import WordParser

        parser = CompositeDocumentParser(
            parsers={
                "application/pdf": PDFParser(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": WordParser(),
                "text/plain": TextParser(),
            },
        )
        path = _create_test_pdf()
        try:
            result = parser.parse(path, "application/pdf")
            assert result.parse_status == "completed"
        finally:
            os.unlink(path)

    def test_route_txt(self) -> None:
        from src.infrastructure.document_parsing.composite_parser import CompositeDocumentParser
        from src.infrastructure.document_parsing.pdf_parser import PDFParser
        from src.infrastructure.document_parsing.text_parser import TextParser
        from src.infrastructure.document_parsing.word_parser import WordParser

        parser = CompositeDocumentParser(
            parsers={
                "application/pdf": PDFParser(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": WordParser(),
                "text/plain": TextParser(),
            },
        )
        path = _create_test_txt()
        try:
            result = parser.parse(path, "text/plain")
            assert result.parse_status == "completed"
        finally:
            os.unlink(path)

    def test_route_unknown_returns_failed(self) -> None:
        from src.infrastructure.document_parsing.composite_parser import CompositeDocumentParser
        from src.infrastructure.document_parsing.pdf_parser import PDFParser
        from src.infrastructure.document_parsing.text_parser import TextParser
        from src.infrastructure.document_parsing.word_parser import WordParser

        parser = CompositeDocumentParser(
            parsers={
                "application/pdf": PDFParser(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": WordParser(),
                "text/plain": TextParser(),
            },
        )
        path = _create_test_txt()
        try:
            result = parser.parse(path, "application/unknown")
            assert result.parse_status == "failed"
            assert result.error_message is not None
            assert "不支持的 MIME" in result.error_message
        finally:
            os.unlink(path)


# ===================================================================
# Service 层集成测试：通过 DocumentParsingService 完整流水线
# 使用真实 PostgreSQL + 真实 MinIO + 真实解析器 + 真实 EventBus
# ===================================================================


class TestDocumentParsingServicePipeline:
    """测试 DocumentParsingService 完整编排流水线

    真实组件：PostgreSQL（schema 隔离）、MinIO、解析器、DualChannelEventBus（Redis + Outbox → RabbitMQ）。
    """

    def _make_composite_parser(self):
        from src.infrastructure.document_parsing.composite_parser import CompositeDocumentParser
        from src.infrastructure.document_parsing.pdf_parser import PDFParser
        from src.infrastructure.document_parsing.text_parser import TextParser
        from src.infrastructure.document_parsing.word_parser import WordParser

        return CompositeDocumentParser(
            parsers={
                "application/pdf": PDFParser(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": WordParser(),
                "text/plain": TextParser(),
            },
        )

    @pytest.mark.asyncio
    async def test_parse_pdf_through_service(
        self,
        pg_session: AsyncSession,
        repo: PostgreSQLDocumentRepository,
        minio_storage,
        event_bus,
    ) -> None:
        """完整流水线：PDF 文件上传到 MinIO → 通过 Service 下载并解析"""
        from src.application.services.document_parsing_service import DocumentParsingService

        # 创建 PDF 并上传到 MinIO
        pdf_path = _create_test_pdf_with_text("Service Pipeline Test PDF")
        object_key = f"test-{uuid.uuid4().hex[:8]}.pdf"
        try:
            await _upload_to_minio_async(minio_storage, "raw-documents", object_key, pdf_path)

            parser = self._make_composite_parser()
            service = DocumentParsingService(
                document_repository=repo,
                document_storage=minio_storage,
                event_publisher=event_bus,
                document_parser=parser,
            )

            doc_id = uuid.uuid4()
            tenant_id = "test-tenant"
            doc = Document(
                document_id=doc_id,
                filename="test.pdf",
                mime_type="application/pdf",
                tenant_id=tenant_id,
            )
            doc.metadata["storage_object_key"] = object_key

            await repo.save(doc)
            await pg_session.flush()

            result = await service.parse_document(doc_id, tenant_id)

            assert result.parse_status == ParseStatus.COMPLETED
            parse_result = result.metadata.get("parse_result")
            assert parse_result is not None, f"metadata 应包含 parse_result，实际: {result.metadata}"
            assert parse_result["parse_status"] == "completed"
            assert len(parse_result["pages"]) >= 1

            # 验证事件已写入 Outbox（RELIABLE 通道）
            outbox_result = await pg_session.execute(
                sa_select(OutboxModel).where(OutboxModel.event_type == "DocumentProcessed")
            )
            outbox_entries = list(outbox_result.scalars().all())
            assert len(outbox_entries) == 1, f"Outbox 应有 1 条 DocumentProcessed 事件，实际: {len(outbox_entries)}"
            assert outbox_entries[0].status == "pending"
        finally:
            _cleanup(pdf_path)

    @pytest.mark.asyncio
    async def test_parse_docx_through_service(
        self,
        pg_session: AsyncSession,
        repo: PostgreSQLDocumentRepository,
        minio_storage,
        event_bus,
    ) -> None:
        """完整流水线：DOCX 文件上传到 MinIO → 通过 Service 下载并解析"""
        from src.application.services.document_parsing_service import DocumentParsingService

        docx_path = _create_test_docx()
        object_key = f"test-{uuid.uuid4().hex[:8]}.docx"
        try:
            await _upload_to_minio_async(minio_storage, "raw-documents", object_key, docx_path)

            parser = self._make_composite_parser()
            service = DocumentParsingService(
                document_repository=repo,
                document_storage=minio_storage,
                event_publisher=event_bus,
                document_parser=parser,
            )

            doc_id = uuid.uuid4()
            tenant_id = "test-tenant"
            doc = Document(
                document_id=doc_id,
                filename="test.docx",
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                tenant_id=tenant_id,
            )
            doc.metadata["storage_object_key"] = object_key

            await repo.save(doc)
            await pg_session.flush()

            result = await service.parse_document(doc_id, tenant_id)

            assert result.parse_status == ParseStatus.COMPLETED
            parse_result = result.metadata.get("parse_result")
            assert parse_result is not None
            all_text = " ".join(t["content"] for p in parse_result["pages"] for t in p["texts"])
            assert "集成测试段落" in all_text

            # 验证事件已写入 Outbox
            outbox_result = await pg_session.execute(
                sa_select(OutboxModel).where(OutboxModel.event_type == "DocumentProcessed")
            )
            outbox_entries = list(outbox_result.scalars().all())
            assert len(outbox_entries) == 1
        finally:
            _cleanup(docx_path)

    @pytest.mark.asyncio
    async def test_parse_txt_through_service(
        self,
        pg_session: AsyncSession,
        repo: PostgreSQLDocumentRepository,
        minio_storage,
        event_bus,
    ) -> None:
        """完整流水线：TXT 文件上传到 MinIO → 通过 Service 下载并解析"""
        from src.application.services.document_parsing_service import DocumentParsingService

        txt_path = _create_test_txt("utf-8", "Service流水线测试\n\n第二段内容")
        object_key = f"test-{uuid.uuid4().hex[:8]}.txt"
        try:
            await _upload_to_minio_async(minio_storage, "raw-documents", object_key, txt_path)

            parser = self._make_composite_parser()
            service = DocumentParsingService(
                document_repository=repo,
                document_storage=minio_storage,
                event_publisher=event_bus,
                document_parser=parser,
            )

            doc_id = uuid.uuid4()
            tenant_id = "test-tenant"
            doc = Document(
                document_id=doc_id,
                filename="test.txt",
                mime_type="text/plain",
                tenant_id=tenant_id,
            )
            doc.metadata["storage_object_key"] = object_key

            await repo.save(doc)
            await pg_session.flush()

            result = await service.parse_document(doc_id, tenant_id)

            assert result.parse_status == ParseStatus.COMPLETED
            parse_result = result.metadata.get("parse_result")
            assert parse_result is not None
            all_text = " ".join(t["content"] for p in parse_result["pages"] for t in p["texts"])
            assert "Service流水线测试" in all_text

            # 验证事件已写入 Outbox
            outbox_result = await pg_session.execute(
                sa_select(OutboxModel).where(OutboxModel.event_type == "DocumentProcessed")
            )
            outbox_entries = list(outbox_result.scalars().all())
            assert len(outbox_entries) == 1
        finally:
            _cleanup(txt_path)

    @pytest.mark.asyncio
    async def test_document_not_found_returns_failed(
        self,
        pg_session: AsyncSession,
        repo: PostgreSQLDocumentRepository,
        minio_storage,
        event_bus,
    ) -> None:
        """文档不存在时返回 FAILED 状态，不发布事件"""
        from src.application.services.document_parsing_service import DocumentParsingService

        parser = self._make_composite_parser()
        service = DocumentParsingService(
            document_repository=repo,
            document_storage=minio_storage,
            event_publisher=event_bus,
            document_parser=parser,
        )

        result = await service.parse_document(uuid.uuid4(), "nonexistent")

        assert result.parse_status == ParseStatus.FAILED
        assert result.metadata.get("parse_error") == "文档不存在"

        # 验证未写入 Outbox
        outbox_result = await pg_session.execute(sa_select(OutboxModel).where(OutboxModel.event_type == "DocumentProcessed"))
        outbox_entries = list(outbox_result.scalars().all())
        assert len(outbox_entries) == 0

    @pytest.mark.asyncio
    async def test_status_update_persisted_to_db(
        self,
        pg_session: AsyncSession,
        repo: PostgreSQLDocumentRepository,
        minio_storage,
        event_bus,
    ) -> None:
        """验证解析后状态变更已持久化到 PostgreSQL"""
        from src.application.services.document_parsing_service import DocumentParsingService

        pdf_path = _create_test_pdf_with_text("Status Persistence Test")
        object_key = f"test-{uuid.uuid4().hex[:8]}.pdf"
        try:
            await _upload_to_minio_async(minio_storage, "raw-documents", object_key, pdf_path)

            parser = self._make_composite_parser()
            service = DocumentParsingService(
                document_repository=repo,
                document_storage=minio_storage,
                event_publisher=event_bus,
                document_parser=parser,
            )

            doc_id = uuid.uuid4()
            tenant_id = "test-tenant"
            doc = Document(
                document_id=doc_id,
                filename="status.pdf",
                mime_type="application/pdf",
                tenant_id=tenant_id,
            )
            doc.metadata["storage_object_key"] = object_key

            await repo.save(doc)
            await pg_session.flush()

            await service.parse_document(doc_id, tenant_id)

            # 从 DB 重新加载验证持久化
            from src.domain.ports.document_repository import DocumentQuery

            reloaded = await repo.find(DocumentQuery(tenant_id=tenant_id, document_id=doc_id))
            assert reloaded is not None
            assert reloaded.parse_status == ParseStatus.COMPLETED
            assert "parse_result" in reloaded.metadata

            # 验证事件已写入 Outbox
            outbox_result = await pg_session.execute(
                sa_select(OutboxModel).where(OutboxModel.event_type == "DocumentProcessed")
            )
            outbox_entries = list(outbox_result.scalars().all())
            assert len(outbox_entries) == 1
        finally:
            _cleanup(pdf_path)


# ===================================================================
# 并发解析测试（AC-6: ≥10 并发）
# ===================================================================


class TestConcurrentParsing:
    """并发解析测试"""

    def test_concurrent_parse_10_documents(self) -> None:
        from src.infrastructure.document_parsing.pdf_parser import PDFParser

        parser = PDFParser()
        paths = [_create_test_pdf() for _ in range(10)]
        try:

            async def parse_all() -> list[ParsedDocument]:
                tasks = [asyncio.to_thread(parser.parse, p, "application/pdf") for p in paths]
                return await asyncio.gather(*tasks)

            loop = asyncio.new_event_loop()
            try:
                results = loop.run_until_complete(parse_all())
            finally:
                loop.close()

            assert len(results) == 10
            for result in results:
                assert result.parse_status == "completed"
        finally:
            for p in paths:
                os.unlink(p)
