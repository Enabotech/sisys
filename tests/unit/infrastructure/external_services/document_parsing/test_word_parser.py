"""Word 文档解析器单元测试

TDD 红阶段：测试 WordParser 的文本提取、表格提取、段落样式、DOC 格式拒绝。
使用 python-docx 创建 fixture DOCX 文件。
"""

from __future__ import annotations

import os
import tempfile

from docx import Document as DocxDocument


def _create_docx_with_paragraphs(texts: list[str]) -> str:
    """创建包含段落文本的 DOCX fixture"""
    doc = DocxDocument()
    for text in texts:
        doc.add_paragraph(text)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def _create_docx_with_table(rows: list[list[str]]) -> str:
    """创建包含表格的 DOCX fixture"""
    doc = DocxDocument()
    table = doc.add_table(rows=len(rows), cols=len(rows[0]) if rows else 0)
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            table.rows[i].cells[j].text = cell_text
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def _create_docx_with_heading(heading_text: str, body_text: str) -> str:
    """创建包含标题和正文的 DOCX fixture"""
    doc = DocxDocument()
    doc.add_heading(heading_text, level=1)
    doc.add_paragraph(body_text)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def _create_empty_docx() -> str:
    """创建空 DOCX"""
    doc = DocxDocument()
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


class TestWordParserCreation:
    """WordParser 构造测试"""

    def test_create_parser(self) -> None:
        from src.infrastructure.external_services.document_parsing.word_parser import WordParser

        parser = WordParser()
        assert parser is not None


class TestWordParserTextExtraction:
    """文本提取测试"""

    def test_extract_single_paragraph(self) -> None:
        from src.infrastructure.external_services.document_parsing.word_parser import WordParser

        parser = WordParser()
        path = _create_docx_with_paragraphs(["Hello World"])
        try:
            result = parser.parse(path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            assert result.parse_status == "completed"
            assert len(result.pages) >= 1
            all_text = " ".join(t.content for p in result.pages for t in p.texts)
            assert "Hello World" in all_text
        finally:
            os.unlink(path)

    def test_extract_multiple_paragraphs(self) -> None:
        from src.infrastructure.external_services.document_parsing.word_parser import WordParser

        parser = WordParser()
        path = _create_docx_with_paragraphs(["第一段", "第二段", "第三段"])
        try:
            result = parser.parse(path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            assert result.parse_status == "completed"
            all_text = " ".join(t.content for p in result.pages for t in p.texts)
            assert "第一段" in all_text
            assert "第三段" in all_text
        finally:
            os.unlink(path)


class TestWordParserTableExtraction:
    """表格提取测试"""

    def test_extract_table(self) -> None:
        from src.infrastructure.external_services.document_parsing.word_parser import WordParser

        parser = WordParser()
        path = _create_docx_with_table([["Name", "Age"], ["Alice", "30"]])
        try:
            result = parser.parse(path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            assert result.parse_status == "completed"
            assert len(result.pages) >= 1
            # 应该包含至少一个表格
            all_tables = [t for p in result.pages for t in p.tables]
            assert len(all_tables) >= 1
            assert all_tables[0].rows[0] == ["Name", "Age"]
            assert all_tables[0].rows[1] == ["Alice", "30"]
        finally:
            os.unlink(path)


class TestWordParserHeading:
    """标题样式测试"""

    def test_heading_extracted(self) -> None:
        from src.infrastructure.external_services.document_parsing.word_parser import WordParser

        parser = WordParser()
        path = _create_docx_with_heading("战略规划", "这是正文内容")
        try:
            result = parser.parse(path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            assert result.parse_status == "completed"
            all_text = " ".join(t.content for p in result.pages for t in p.texts)
            assert "战略规划" in all_text
            assert "正文内容" in all_text
        finally:
            os.unlink(path)


class TestWordParserEdgeCases:
    """边界场景测试"""

    def test_empty_docx(self) -> None:
        """空 DOCX（无段落无表格）应返回 failed（AC-2 要求）"""
        from src.infrastructure.external_services.document_parsing.word_parser import WordParser

        parser = WordParser()
        path = _create_empty_docx()
        try:
            result = parser.parse(path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            assert result.is_failed()
            assert result.error_message is not None
            assert "空" in result.error_message
        finally:
            os.unlink(path)

    def test_doc_format_rejected(self) -> None:
        """验证旧版 DOC 格式返回 failed"""
        from src.infrastructure.external_services.document_parsing.word_parser import WordParser

        parser = WordParser()
        # 创建一个非 DOCX 格式的文件（模拟 DOC）
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".doc")
        tmp.write(b"not a valid docx")
        tmp.close()
        try:
            result = parser.parse(tmp.name, "application/msword")
            assert result.parse_status == "failed"
            assert result.error_message is not None
            assert "DOCX" in result.error_message or "docx" in result.error_message.lower()
        finally:
            os.unlink(tmp.name)

    def test_output_to_dict_json_serializable(self) -> None:
        import json

        from src.infrastructure.external_services.document_parsing.word_parser import WordParser

        parser = WordParser()
        path = _create_docx_with_paragraphs(["test"])
        try:
            result = parser.parse(path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            d = result.to_dict()
            json_str = json.dumps(d, ensure_ascii=False)
            assert len(json_str) > 0
        finally:
            os.unlink(path)


class TestWordParserSizeLimit:
    """DOCX 文件大小上限保护测试（防御内嵌 OOXML 解压炸弹）"""

    def test_oversized_docx_returns_failed(self, monkeypatch) -> None:
        """超过 MAX_DOCX_BYTES 应返回 failed"""
        from src.infrastructure.external_services.document_parsing import _limits
        from src.infrastructure.external_services.document_parsing.word_parser import WordParser

        monkeypatch.setattr(
            "os.path.getsize",
            lambda _path: _limits.MAX_DOCX_BYTES + 1,
        )
        parser = WordParser()
        result = parser.parse("/tmp/whatever.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert result.is_failed()
        assert "50MB" in (result.error_message or "")

    def test_getsize_oserror_returns_failed(self, monkeypatch) -> None:
        """os.path.getsize 抛出 OSError 时应返回 failed 而非异常穿透"""
        from src.infrastructure.external_services.document_parsing.word_parser import WordParser

        monkeypatch.setattr("os.path.getsize", lambda _: (_ for _ in ()).throw(OSError("Permission denied")))
        parser = WordParser()
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        result = parser.parse("/inaccessible/file.docx", mime)
        assert result.is_failed()
        assert "权限" in (result.error_message or "")


class TestWordParserExceptionSanitization:
    """DOCX 异常信息脱敏测试"""

    def test_corrupt_docx_returns_failed_without_leaking_path(self) -> None:
        """损坏 DOCX 应返回 failed 且 error_message 不含路径"""
        from src.infrastructure.external_services.document_parsing.word_parser import WordParser

        parser = WordParser()
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
            f.write(b"this is not a real docx zip content")
            path = f.name
        try:
            result = parser.parse(path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            assert result.is_failed()
            assert result.error_message is not None
            assert path not in result.error_message
            assert "this is not a real docx" not in result.error_message
        finally:
            os.unlink(path)
